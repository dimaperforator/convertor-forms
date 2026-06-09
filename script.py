import json
import os
import sys
import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.style import WD_STYLE_TYPE
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.pagesizes import A5
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

def resource(name):
    if getattr(sys, "frozen", False):
        return os.path.join(os.path.dirname(sys.executable), name)
    return os.path.join(os.path.dirname(__file__), name)

FONT_PATH = resource("DejaVuSans.ttf")

if not os.path.exists(FONT_PATH):
    messagebox.showerror("Ошибка", "Нет файла DejaVuSans.ttf")
    sys.exit()

pdfmetrics.registerFont(TTFont("DejaVu", FONT_PATH))

SKIP_FIELDS = [
    "ID",
    "Время создания",
    "Я даю свое согласие на обработку персональных данных"
]

def apply_style(text, style):
    if "bold" in style:
        text = f"<b>{text}</b>"
    if "italic" in style:
        text = f"<i>{text}</i>"
    if "underline" in style:
        text = f"<u>{text}</u>"
    return text

def make_pdf(data, out, q_style, a_style, title_style):
    doc = SimpleDocTemplate(out, pagesize=A5, leftMargin=20, rightMargin=20, topMargin=20, bottomMargin=20)
    story = []

    for i, form in enumerate(data, 1):
        story.append(
            Paragraph(
                apply_style(f"Заявка №{i}", title_style),
                ParagraphStyle("title", fontName="DejaVu", fontSize=16, leading=20)
            )
        )

        story.append(Spacer(1, 12))

        for q, a in form:
            q_text = apply_style(f"{q}:", q_style)
            a_text = a if isinstance(a, str) and a.startswith("http") else apply_style(str(a), a_style)

            story.append(
                Paragraph(
                    q_text + " " + a_text,
                    ParagraphStyle("line", fontName="DejaVu", fontSize=11, leading=15)
                )
            )

            story.append(Spacer(1, 10))

        if i != len(data):
            story.append(PageBreak())

    doc.build(story)

def make_word(data, out, q_style, a_style, title_style):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(148)
    section.page_height = Mm(210)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    styles = doc.styles

    if "TitleStyle" not in styles:
        title = styles.add_style("TitleStyle", WD_STYLE_TYPE.PARAGRAPH)
        title.font.name = "Times New Roman"
        title.font.size = Pt(16)
        title.font.bold = "bold" in title_style
        title.font.italic = "italic" in title_style

    if "Question" not in styles:
        qst = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
        qst.font.name = "Times New Roman"
        qst.font.size = Pt(12)
        qst.font.bold = "bold" in q_style
        qst.font.italic = "italic" in q_style

    if "Answer" not in styles:
        ans = styles.add_style("Answer", WD_STYLE_TYPE.PARAGRAPH)
        ans.font.name = "Times New Roman"
        ans.font.size = Pt(12)
        ans.font.bold = "bold" in a_style
        ans.font.italic = "italic" in a_style

    for i, form in enumerate(data, 1):
        doc.add_paragraph(f"Заявка №{i}", style="TitleStyle")

        for q, a in form:
            doc.add_paragraph(f"{q}:", style="Question")
            doc.add_paragraph(str(a), style="Answer")

        if i != len(data):
            doc.add_page_break()

    doc.save(out)

class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Конвертер анкет")
        self.geometry("700x1000")

        self.json_path = ""
        self.loaded_data = []
        self.field_vars = {}

        self.main_frame = ctk.CTkScrollableFrame(self)
        self.main_frame.pack(fill="both", expand=True)

        ctk.CTkLabel(self.main_frame, text="Конвертер анкет", font=("Arial", 32, "bold")).pack(pady=25)

        self.json_btn = ctk.CTkButton(
            self.main_frame,
            text="Выбрать JSON",
            width=260,
            height=55,
            command=self.select_json
        )
        self.json_btn.pack(pady=12)

        self.path_label = ctk.CTkLabel(self.main_frame, text="Файл не выбран")
        self.path_label.pack()

        ctk.CTkLabel(self.main_frame, text="Поля для выгрузки", font=("Arial", 22, "bold")).pack(pady=(20, 10))

        self.fields_frame = ctk.CTkFrame(self.main_frame)
        self.fields_frame.pack(fill="x", padx=20)

        self.skip_empty = ctk.BooleanVar(value=False)

        ctk.CTkCheckBox(
            self.main_frame,
            text="Пропускать пустые ответы",
            variable=self.skip_empty
        ).pack(pady=10)

        self.format_var = ctk.StringVar(value="word")

        ctk.CTkLabel(self.main_frame, text="Формат файла", font=("Arial", 22, "bold")).pack(pady=(20, 10))

        ctk.CTkRadioButton(self.main_frame, text="Word", variable=self.format_var, value="word").pack()
        ctk.CTkRadioButton(self.main_frame, text="PDF", variable=self.format_var, value="pdf").pack()

        ctk.CTkLabel(self.main_frame, text="Стили", font=("Arial", 22, "bold")).pack(pady=(20, 10))

        self.title_bold = ctk.BooleanVar()
        self.title_italic = ctk.BooleanVar()
        self.q_bold = ctk.BooleanVar()
        self.q_italic = ctk.BooleanVar()
        self.a_bold = ctk.BooleanVar()
        self.a_italic = ctk.BooleanVar()

        ctk.CTkLabel(self.main_frame, text="Заявка", font=("Arial", 18, "bold")).pack()
        ctk.CTkCheckBox(self.main_frame, text="Жирный", variable=self.title_bold).pack()
        ctk.CTkCheckBox(self.main_frame, text="Курсив", variable=self.title_italic).pack()

        ctk.CTkLabel(self.main_frame, text="Вопрос", font=("Arial", 18, "bold")).pack(pady=(15, 0))
        ctk.CTkCheckBox(self.main_frame, text="Жирный", variable=self.q_bold).pack()
        ctk.CTkCheckBox(self.main_frame, text="Курсив", variable=self.q_italic).pack()

        ctk.CTkLabel(self.main_frame, text="Ответ", font=("Arial", 18, "bold")).pack(pady=(15, 0))
        ctk.CTkCheckBox(self.main_frame, text="Жирный", variable=self.a_bold).pack()
        ctk.CTkCheckBox(self.main_frame, text="Курсив", variable=self.a_italic).pack()

        buttons_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        buttons_frame.pack(pady=30)

        ctk.CTkButton(buttons_frame, text="Сохранить", width=220, height=55, command=self.create_file).grid(row=0, column=0, padx=10)
        ctk.CTkButton(buttons_frame, text="Отмена", width=180, height=55, command=self.destroy).grid(row=0, column=1, padx=10)

    def select_json(self):
        path = filedialog.askopenfilename(title="Выберите JSON", filetypes=[("JSON files", "*.json")])

        if not path:
            return

        with open(path, "r", encoding="utf-8-sig") as f:
            data = json.load(f)

        self.loaded_data = data
        self.json_path = path
        self.path_label.configure(text=os.path.basename(path))

        for widget in self.fields_frame.winfo_children():
            widget.destroy()

        self.field_vars.clear()

        fields = set()

        for form in data:
            for q, a in form:
                if q not in SKIP_FIELDS:
                    fields.add(q)

        for field in sorted(fields):
            var = ctk.BooleanVar(value=True)
            self.field_vars[field] = var
            ctk.CTkCheckBox(self.fields_frame, text=field, variable=var).pack(anchor="w", padx=10, pady=2)

    def get_output_file(self):
        base = os.path.splitext(os.path.basename(self.json_path))[0]

        if self.format_var.get() == "pdf":
            return filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF", "*.pdf")],
                initialfile=f"{base}.pdf"
            )

        return filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")],
            initialfile=f"{base}.docx"
        )

    def create_file(self):
        if not self.json_path:
            messagebox.showerror("Ошибка", "Выберите JSON файл")
            return

        out = self.get_output_file()

        if not out:
            return

        title_style = []
        q_style = []
        a_style = []

        if self.title_bold.get():
            title_style.append("bold")
        if self.title_italic.get():
            title_style.append("italic")

        if self.q_bold.get():
            q_style.append("bold")
        if self.q_italic.get():
            q_style.append("italic")

        if self.a_bold.get():
            a_style.append("bold")
        if self.a_italic.get():
            a_style.append("italic")

        selected_fields = {k for k, v in self.field_vars.items() if v.get()}

        filtered_data = []

        for form in self.loaded_data:
            row = []

            for q, a in form:
                if q in SKIP_FIELDS:
                    continue

                if q not in selected_fields:
                    continue

                if self.skip_empty.get() and str(a).strip() == "":
                    continue

                row.append((q, a))

            filtered_data.append(row)

        if self.format_var.get() == "pdf":
            make_pdf(filtered_data, out, q_style, a_style, title_style)
        else:
            make_word(filtered_data, out, q_style, a_style, title_style)

        messagebox.showinfo("Готово", f"Файл создан:\n{out}")

App().mainloop()
